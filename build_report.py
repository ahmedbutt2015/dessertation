#!/usr/bin/env python3
"""
Build Supplementary_Report_Ahmed_Fayyaz_Butt.docx
- All prose from SUPPLEMENTARY_DRAFT.md
- Code inline in Chapter 5 subsections (not appendix)
- All images embedded inline
- Three factual fixes applied
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = Path('/Users/ahmedthebutt/Documents/ULSTER/dessertation')
DRAFT = BASE / 'report' / 'SUPPLEMENTARY_DRAFT.md'
OUT   = BASE / 'report' / 'Supplementary_Report_Ahmed_Fayyaz_Butt.docx'
IMGS  = BASE / 'results'

IMG_MAP = {
    'fig1_pipeline_diagram.png':       IMGS / 'fig1_pipeline_diagram.png',
    'fig2_lightcurve_to_gaf.png':      IMGS / 'fig2_lightcurve_to_gaf.png',
    'fig3_four_settings_barchart.png': IMGS / 'fig3_four_settings_barchart.png',
    'fig5_explanation_example.png':    IMGS / 'fig5_explanation_example.png',
    'fig6_gantt_chart.png':            IMGS / 'fig6_gantt_chart.png',
    'fig6_rank_sweep.png':             IMGS / 'fig6_rank_sweep.png',
    'fig7_multiseed_stability.png':    IMGS / 'fig7_multiseed_stability.png',
    'lora_r16_roc_curve.png':          IMGS / 'lora_r16_roc_curve.png',
    'lora_r16_confusion_matrix.png':   IMGS / 'lora_r16_confusion_matrix.png',
    # App screenshots (S9, S11, S12, S13)
    'app/3.png': IMGS / 'app' / '3.png',
    'app/1.png': IMGS / 'app' / '1.png',
    'app/2.png': IMGS / 'app' / '2.png',
    'app/4.png': IMGS / 'app' / '4.png',
}

# ---------------------------------------------------------------------------
# Code listings per section
# ---------------------------------------------------------------------------
CODE_34 = [
("Listing 3.1 — Notebook 01 (reference): Lightkurve MAST acquisition pipeline",
"""import lightkurve as lk
from pyts.image import GramianAngularField

def acquire_koi_gaf(kepid, period, t0, n_bins=1024, gaf_size=64):
    # 1. Download all PDCSAP quarters from NASA MAST
    search = lk.search_lightcurve(f'KIC {kepid}', mission='Kepler',
                                   cadence='long', author='Kepler')
    lcs = search.download_all()
    if lcs is None or len(lcs) == 0:
        raise ValueError(f'No Kepler data for KIC {kepid}')

    # 2. Stitch quarters; normalise; phase-fold on catalogue params
    lc     = lcs.stitch().remove_nans().normalize()
    folded = lc.fold(period=period, epoch_time=t0)
    binned = folded.bin(n_bins=n_bins).flux.value

    # 3. Rescale to [-1, 1] and apply GAF transform
    mn, mx = binned.min(), binned.max()
    if mx - mn < 1e-8:
        raise ValueError(f'Flat-line for KIC {kepid}')
    binned = 2.0 * (binned - mn) / (mx - mn) - 1.0
    gaf    = GramianAngularField(image_size=gaf_size, method='summation')
    return gaf.fit_transform(binned.reshape(1, -1))[0].astype('float32')

# Operational failure at scale:
# ~7,600 KOIs x 17 quarters x ~7 s/file = 15-21 hrs download time
# Kaggle CPU session limit: 9 hours -> session terminated, /kaggle/working wiped
# Contingency activated: Macedo & Zalewski 2024 (Mendeley) -- see Section 3.5"""),
]

CODE_53 = [
("Listing 5.1 — Notebook 01b: Load Mendeley CSV and apply flat-line filter",
"""GLOBAL_PATH = MENDELEY_INPUT / 'all_global.csv'
raw_df    = pd.read_csv(GLOBAL_PATH)
flux_cols = [str(i) for i in range(CFG['phase_bins'])]   # '0' ... '2000'

flux_matrix = raw_df[flux_cols].values.astype(np.float32)  # (5302, 2001)
f_min  = flux_matrix.min(axis=1, keepdims=True)
f_max  = flux_matrix.max(axis=1, keepdims=True)
denom  = f_max - f_min
flat_mask   = denom.squeeze() >= 1e-8       # discard constant-flux rows
flux_matrix = flux_matrix[flat_mask]
flux_scaled = 2.0 * (flux_matrix - f_min[flat_mask]) / denom[flat_mask] - 1.0
labels_raw  = raw_df['label'].values[flat_mask]
print(f"Rows after flat-line filter: {flat_mask.sum()} / {len(raw_df)}")"""),

("Listing 5.2 — Notebook 01b: GAF encoding (64x64 Gramian Angular Field)",
"""from pyts.image import GramianAngularField

gaf_transform = GramianAngularField(image_size=CFG['gaf_size'], method='summation')
X_raw = np.stack([flux_scaled[i] for i in range(len(flux_scaled))])   # (N, 2001)
X     = gaf_transform.fit_transform(X_raw).astype(np.float32)          # (N, 64, 64)
y     = (labels_raw == 'CONFIRMED').astype(np.int64)
print(f"X: {X.shape}  dtype={X.dtype}")
print(f"y: {y.shape}  confirmed={y.mean():.1%}")"""),
]

CODE_54 = [
("Listing 5.3 — Notebook 01b: Stratified 70/15/15 split and save",
"""from sklearn.model_selection import train_test_split

X_tv, X_test, y_tv, y_test = train_test_split(
    X, y, test_size=CFG['test_size'], stratify=y,
    random_state=CFG['random_seed'])
val_frac = CFG['val_size'] / (1.0 - CFG['test_size'])
X_train, X_val, y_train, y_val = train_test_split(
    X_tv, y_tv, test_size=val_frac, stratify=y_tv,
    random_state=CFG['random_seed'])

for name, X_s, y_s in [('train',X_train,y_train),('val',X_val,y_val),('test',X_test,y_test)]:
    print(f"{name:5s}: {len(X_s):4d} samples | confirmed: {y_s.mean():.1%}")

np.savez_compressed(DATASET_PATH, X_train=X_train, y_train=y_train,
                    X_val=X_val, y_val=y_val, X_test=X_test, y_test=y_test)"""),

("Listing 5.4 — Notebook 02: GAFDataset class (resize to 224x224, 3 channels, normalise)",
"""class GAFDataset(Dataset):
    def __init__(self, X, y):
        self.X = torch.from_numpy(X).unsqueeze(1)   # add channel dim
        self.y = torch.from_numpy(y).long()

    def __len__(self):
        return len(self.y)

    def __getitem__(self, idx):
        img = self.X[idx]                                     # (1, 64, 64)
        img = F.interpolate(img.unsqueeze(0), size=224,
                            mode='bilinear',
                            align_corners=False).squeeze(0)   # (1, 224, 224)
        img = img.repeat(3, 1, 1)                             # (3, 224, 224)
        mean = torch.tensor([0.5, 0.5, 0.5]).view(3, 1, 1)
        std  = torch.tensor([0.5, 0.5, 0.5]).view(3, 1, 1)
        return (img - mean) / std, self.y[idx]

# Class weights from training partition counts
n_confirmed   = int(y_train.sum())
n_fp          = len(y_train) - n_confirmed
w_confirmed   = len(y_train) / (2.0 * n_confirmed)
w_fp          = len(y_train) / (2.0 * n_fp)
CLASS_WEIGHTS = torch.tensor([w_fp, w_confirmed], dtype=torch.float32).to(DEVICE)
print(f"n_confirmed={n_confirmed}  n_fp={n_fp}")
print(f"w_confirmed={w_confirmed:.3f}  w_fp={w_fp:.3f}")"""),
]

CODE_55 = [
("Listing 5.5 — Notebook 02: ViT-B/16 factory and evaluation function",
"""import timm, time, torch.nn as nn
from sklearn.metrics import f1_score, roc_auc_score

def make_vit():
    model = timm.create_model('vit_base_patch16_224', pretrained=True, num_classes=0)
    model.head = nn.Linear(768, 2)
    nn.init.xavier_uniform_(model.head.weight)
    nn.init.zeros_(model.head.bias)
    return model

def evaluate(model, loader, device):
    model.eval()
    all_preds, all_probs, all_labels = [], [], []
    t0 = time.time()
    with torch.no_grad():
        for imgs, labels in loader:
            imgs   = imgs.to(device)
            logits = model(imgs)
            probs  = torch.softmax(logits, dim=1)[:, 1]
            preds  = logits.argmax(dim=1)
            all_preds.append(preds.cpu())
            all_probs.append(probs.cpu())
            all_labels.append(labels.cpu())
    elapsed_ms = (time.time() - t0) * 1000
    P = torch.cat(all_preds).numpy()
    Pr= torch.cat(all_probs).numpy()
    L = torch.cat(all_labels).numpy()
    return f1_score(L, P, average='macro'), roc_auc_score(L, Pr), elapsed_ms/len(loader.dataset)

total_params = sum(p.numel() for p in make_vit().parameters())
print(f"Total parameters: {total_params:,}")   # 85,878,786"""),
]

CODE_56 = [
("Listing 5.6 — Notebook 02: Zero-shot, one-shot, and few-shot evaluation",
"""# Zero-shot: no training at all
model_zs = make_vit().to(DEVICE)
f1, auc, ms = evaluate(model_zs, test_loader, DEVICE)
log_result('zero-shot', f1, auc, ms)

def get_n_shot_loader(dataset, y_array, n_per_class, seed=42):
    rng = np.random.RandomState(seed)
    indices = []
    for cls in [0, 1]:
        cls_idx = np.where(y_array == cls)[0]
        indices.extend(rng.choice(cls_idx, n_per_class, replace=False).tolist())
    return DataLoader(Subset(dataset, indices), batch_size=len(indices), shuffle=True)

def train_head_only(model, support_loader, device, epochs=50, lr=1e-3):
    for name, param in model.named_parameters():
        param.requires_grad = name.startswith('head')
    opt = torch.optim.Adam(filter(lambda p: p.requires_grad, model.parameters()), lr=lr)
    crit = nn.CrossEntropyLoss(weight=CLASS_WEIGHTS)
    model.train()
    for _ in range(epochs):
        for imgs, labels in support_loader:
            imgs, labels = imgs.to(device), labels.to(device)
            opt.zero_grad(); crit(model(imgs), labels).backward(); opt.step()
    return model

# One-shot (1 example/class)
model_1s = train_head_only(make_vit().to(DEVICE),
                            get_n_shot_loader(train_ds, y_train, 1), DEVICE)
log_result('one-shot', *evaluate(model_1s, test_loader, DEVICE))

# Few-shot (10 examples/class)
model_fs = train_head_only(make_vit().to(DEVICE),
                            get_n_shot_loader(train_ds, y_train, 10), DEVICE)
log_result('few-shot-10', *evaluate(model_fs, test_loader, DEVICE))"""),
]

CODE_57 = [
("Listing 5.7 — Notebook 02: LoRA fine-tuning with PEFT (rank sweep r=4,8,16)",
"""from peft import LoraConfig, get_peft_model

def train_lora(rank, train_loader, val_loader, device, epochs=10, lr=2e-4):
    base     = make_vit()
    lora_cfg = LoraConfig(
        r=rank, lora_alpha=rank*2, target_modules=['qkv'],
        lora_dropout=0.1, bias='none', modules_to_save=['head'])
    model = get_peft_model(base, lora_cfg).to(device)
    model.print_trainable_parameters()

    opt  = torch.optim.AdamW(model.parameters(), lr=lr, weight_decay=1e-2)
    sched= torch.optim.lr_scheduler.CosineAnnealingLR(opt, T_max=epochs)
    crit = nn.CrossEntropyLoss(weight=CLASS_WEIGHTS)

    best_val_f1, best_state = 0.0, None
    for epoch in range(1, epochs+1):
        model.train()
        for imgs, labels in train_loader:
            imgs, labels = imgs.to(device), labels.to(device)
            opt.zero_grad(); loss = crit(model(imgs), labels)
            loss.backward()
            torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)
            opt.step()
        sched.step()
        val_f1, _, _ = evaluate(model, val_loader, device)
        if val_f1 > best_val_f1:
            best_val_f1 = val_f1
            best_state  = {k: v.clone() for k, v in model.state_dict().items()}
    model.load_state_dict(best_state)
    return model

for rank in CFG['lora_ranks']:   # [4, 8, 16]
    model = train_lora(rank, train_loader, val_loader, DEVICE)
    log_result(f'lora-r{rank}', *evaluate(model, test_loader, DEVICE))
    lora_models[rank] = model"""),

("Listing 5.8 — Notebook 06: Multi-seed replication at LoRA r=16",
"""SEEDS = [42, 123, 2026]
rows  = []
for seed in SEEDS:
    torch.manual_seed(seed); np.random.seed(seed); torch.cuda.manual_seed_all(seed)
    model  = train_one_seed(seed)     # same train_lora call with that seed's DataLoader
    f1, auc = evaluate(model, test_loader)
    rows.append({'seed': seed, 'f1_macro': round(f1,4), 'auc_roc': round(auc,4)})
    print(f"Seed {seed}: F1={f1:.4f}  AUC={auc:.4f}")

df = pd.DataFrame(rows)
print(f"Mean F1={df.f1_macro.mean():.3f}+/-{df.f1_macro.std():.3f}  "
      f"AUC={df.auc_roc.mean():.3f}+/-{df.auc_roc.std():.3f}")
df.to_csv(RESULTS_DIR / 'lora_r16_multiseed.csv', index=False)"""),
]

CODE_58 = [
("Listing 5.9 — Notebook 03: Build FAISS cosine-similarity retrieval index",
"""import faiss, pickle
from sklearn.preprocessing import StandardScaler

FEATURES = ['koi_period','koi_duration','koi_depth','koi_prad','koi_steff','koi_srad']

df        = pd.read_csv(koi_files[0], comment='#')
confirmed = df[df['koi_disposition']=='CONFIRMED'].dropna(subset=FEATURES).copy()
print(f"Confirmed with complete 6-param vectors: {len(confirmed)}")

scaler = StandardScaler()
X      = scaler.fit_transform(confirmed[FEATURES].values.astype(np.float32))
faiss.normalize_L2(X)                     # L2-norm: inner product = cosine similarity

index = faiss.IndexFlatIP(X.shape[1])     # exact brute-force search (sub-ms at 2,745 vectors)
index.add(X)
print(f"FAISS index: {index.ntotal} vectors, {X.shape[1]} dims, IndexFlatIP")

# Sanity check: K00001.01 (canonical hot Jupiter) -> 5 short-period giants
D, I = index.search(X[0:1], k=6)
for rank_i,(sim,idx) in enumerate(zip(D[0][1:],I[0][1:]),1):
    r = confirmed.iloc[idx]
    print(f"  {rank_i}. {r.get('kepoi_name','—'):12s}  sim={sim:.4f}  period={r['koi_period']:.1f}d")

faiss.write_index(index, str(OUT_DIR/'faiss_index.bin'))
confirmed.to_csv(OUT_DIR/'confirmed_planets.csv', index=True)
with open(OUT_DIR/'scaler.pkl','wb') as f: pickle.dump(scaler, f)
print("Saved: faiss_index.bin, confirmed_planets.csv, scaler.pkl")"""),
]

CODE_59 = [
("Listing 5.10 — LangGraph typed state and agentic pipeline definition",
"""from typing import TypedDict
from langgraph.graph import StateGraph, END

class ExplanationState(TypedDict):
    koi_id:          str    # KOI identifier
    label:           str    # 'CONFIRMED' or 'FALSE POSITIVE'
    confidence:      float  # P(CONFIRMED) from softmax
    neighbours:      list   # k retrieved neighbour dicts
    context:         str    # formatted evidence string (build_context output)
    explanation:     str    # generated prose (generate_explanation output)
    generation_path: str    # 'llm' or 'template'

def build_context(state: ExplanationState) -> ExplanationState:
    lines = [f"Top {len(state['neighbours'])} confirmed analogues for {state['koi_id']}:"]
    for n in state['neighbours']:
        lines.append(f"  {n['name']}: period={n['koi_period']:.2f}d, "
                     f"depth={n['koi_depth']:.0f}ppm, prad={n['koi_prad']:.2f}Re, "
                     f"steff={n['koi_steff']:.0f}K, sim={n['score']:.4f}")
    state['context'] = "\\n".join(lines)
    return state

def generate_explanation(state: ExplanationState) -> ExplanationState:
    for model_id in LLM_CASCADE:   # Qwen2.5-7B -> Llama-3.2-3B -> Zephyr-7B
        try:
            state['explanation']     = call_llm(model_id, state)
            state['generation_path'] = 'llm'
            return state
        except Exception:
            continue
    state['explanation']     = template_explanation(state)
    state['generation_path'] = 'template'
    return state

# Directed graph: no reverse edges -> strict downstream property (NFR5)
graph = StateGraph(ExplanationState)
graph.add_node("build_context",        build_context)
graph.add_node("generate_explanation", generate_explanation)
graph.add_edge("build_context", "generate_explanation")
graph.add_edge("generate_explanation", END)
graph.set_entry_point("build_context")
pipeline = graph.compile()"""),
]

CODE_510 = [
("Listing 5.11 — app/server.py: Flask REST API (classify and search endpoints)",
"""from flask import Flask, jsonify, request
import numpy as np

app = Flask(__name__)

@app.route("/api/classify", methods=["POST"])
def api_classify():
    data   = request.get_json(silent=True) or {}
    koi_id = str(data.get("koi_id", "")).strip()
    if not koi_id:
        return jsonify({"error": "No KOI ID provided"}), 400

    try:
        koi_info, similar = retrieve_any_koi(koi_id, k=5)
    except ValueError as e:
        return jsonify({"error": str(e)}), 404

    # Classify with saved LoRA r=16 adapter
    if classifier_available():
        gaf = _load_gaf(koi_id)          # deterministic representative test-set image
        label, confidence = classify(gaf)
    else:
        label, confidence = koi_info.get("koi_disposition", "UNKNOWN"), 0.75

    # Generate grounded natural-language explanation
    explanation_text = explain(koi_id, label, confidence, similar)

    return jsonify({
        "koi_id":          koi_id,
        "label":           label,
        "confidence":      round(confidence * 100),
        "nasa_disp":       koi_info.get("koi_disposition", "UNKNOWN"),
        "similar_planets": format_planets(similar),
        "explanation":     explanation_text,
    })

def _load_gaf(koi_id: str) -> np.ndarray:
    # Mendeley dataset has no KOI-name index; use hash for deterministic selection
    if GAF_NPZ.exists():
        data = np.load(GAF_NPZ)
        return data["X_test"][hash(koi_id) % len(data["X_test"])]
    return np.zeros((64, 64), dtype=np.float32)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=7861, debug=False)"""),
]

CODE_511 = [
("Listing 5.12 — Notebook 05: TESS light curve download and GAF preprocessing",
"""import lightkurve as lk
from pyts.image import GramianAngularField

gaf_transform = GramianAngularField(image_size=64, method='summation')

def download_tess_gaf(tic_id, period, t0):
    search = lk.search_lightcurve(f'TIC {tic_id}', mission='TESS', exptime=120)
    if len(search) == 0:
        search = lk.search_lightcurve(f'TIC {tic_id}', mission='TESS')
    lc     = search[0].download().remove_nans().normalize()
    folded = lc.fold(period=period, epoch_time=t0)
    binned = folded.bin(n_bins=2001).flux.value
    binned = np.nan_to_num(binned, nan=np.nanmedian(binned))
    mn, mx = binned.min(), binned.max()
    if mx - mn > 1e-8:
        binned = 2.0 * (binned - mn) / (mx - mn) - 1.0
    return gaf_transform.fit_transform(binned.reshape(1,-1))[0].astype(np.float32)"""),

("Listing 5.13 — Notebook 05: TESS zero-shot inference (Kepler-trained LoRA, no retraining)",
"""from peft import PeftModel

base  = timm.create_model('vit_base_patch16_224', pretrained=True, num_classes=2)
model = PeftModel.from_pretrained(base, str(ADAPTER_PATH)).to(DEVICE)
model.eval()
print("LoRA r=16 loaded -- zero-shot transfer from Kepler to TESS, no retraining")

all_preds, all_probs = [], []
with torch.no_grad():
    for gaf in X_tess:
        t = torch.from_numpy(gaf).unsqueeze(0).unsqueeze(0)
        t = F.interpolate(t, 224, mode='bilinear', align_corners=False)
        t = ((t.repeat(1,3,1,1) - 0.5) / 0.5).to(DEVICE)
        logits = model(t)
        all_probs.append(torch.softmax(logits,1)[0,1].item())
        all_preds.append(int(logits.argmax(1).item()))

f1  = f1_score(y_tess, all_preds, average='macro')
auc = roc_auc_score(y_tess, all_probs)
print(f"TESS zero-shot:  F1={f1:.3f}  AUC={auc:.3f}")
print(f"Kepler test set: F1=0.834  AUC=0.916  (training domain reference)")"""),
]

CODE_INJECTIONS = {
    'ch3-4':  CODE_34,
    'ch5-3':  CODE_53,
    'ch5-4':  CODE_54,
    'ch5-5':  CODE_55,
    'ch5-6':  CODE_56,
    'ch5-7':  CODE_57,
    'ch5-8':  CODE_58,
    'ch5-9':  CODE_59,
    'ch5-10': CODE_510,
    'ch5-11': CODE_511,
}

CODE_TRIGGER = {
    'ch3-5':  'ch3-4',
    'ch5-4':  'ch5-3',
    'ch5-5':  'ch5-4',
    'ch5-6':  'ch5-5',
    'ch5-7':  'ch5-6',
    'ch5-8':  'ch5-7',
    'ch5-9':  'ch5-8',
    'ch5-10': 'ch5-9',
    'ch5-11': 'ch5-10',
    'ch6':    'ch5-11',
}

TEXT_FIXES = [
    ("monitored approximately 150,000 stars continuously",
     "monitored approximately 200,000 stars continuously"),
    ("monitoring approximately 150,000 solar-type stars every 30 minutes",
     "monitoring approximately 200,000 solar-type stars every 30 minutes"),
    ("collected brightness measurements for approximately 150,000 solar-type stars",
     "collected brightness measurements for approximately 200,000 solar-type stars"),
    ("approximately 150,000 solar-type stars every 30 minutes",
     "approximately 200,000 solar-type stars every 30 minutes"),
    ("(1,120 pos vs 2,591 neg)",
     "(~1,537 CONFIRMED, ~2,174 FALSE POSITIVE)"),
    ("(1,120 CONFIRMED, 2,591 FALSE POSITIVE in train)",
     "(~1,537 CONFIRMED, ~2,174 FALSE POSITIVE in train)"),
    ("weight_confirmed = 2.313. If global",
     "weight_confirmed ≈ 1.414. If global"),
    ("weight_confirmed = 2.313",
     "weight_confirmed ≈ 1.414"),
    ("[Houlsby et al., 2019]",  "[S5]"),
    ("Houlsby et al., 2019",    "Houlsby et al. [S5] 2019"),
    ("[Li and Liang, 2021; Lester et al., 2021]", "[S6][S7]"),
    ("[Ben-Zaken et al., 2022]", "[S8]"),
    ("Ben-Zaken et al., 2022",   "Ben-Zaken et al. [S8] 2022"),
    ("Aghajanyan et al. demonstrated empirically",
     "Aghajanyan et al. [S9] demonstrated empirically"),
]

EXTRA_REFS = """
[S5] N. Houlsby et al., "Parameter-Efficient Transfer Learning for NLP," in Proc. 36th International Conference on Machine Learning (ICML), 2019.

[S6] X. L. Li and P. Liang, "Prefix-Tuning: Optimizing Continuous Prompts for Generation," in Proc. 59th Annual Meeting of the Association for Computational Linguistics (ACL), 2021, pp. 4582-4597.

[S7] B. Lester, R. Al-Rfou, and N. Constant, "The Power of Scale for Parameter-Efficient Prompt Tuning," in Proc. Conference on Empirical Methods in Natural Language Processing (EMNLP), 2021, pp. 3045-3059.

[S8] E. Ben-Zaken, S. Ravfogel, and Y. Goldberg, "BitFit: Simple Parameter-Efficient Fine-Tuning for Transformer-Based Masked Language-Models," in Proc. 60th Annual Meeting of the Association for Computational Linguistics (ACL), 2022, pp. 1-9.

[S9] A. Aghajanyan, S. Zettlemoyer, and S. Gupta, "Intrinsic Dimensionality Explains the Effectiveness of Language Model Fine-Tuning," in Proc. 59th Annual Meeting of the Association for Computational Linguistics (ACL), 2021, pp. 7319-7328.
"""

# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def shade_paragraph(para, fill="F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)


def add_code_block(doc, title, code):
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(8)
    tp.paragraph_format.space_after  = Pt(2)
    r = tp.add_run(title)
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'

    for line in code.split('\n'):
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after  = Pt(0)
        cp.paragraph_format.left_indent  = Cm(0.3)
        shade_paragraph(cp)
        r = cp.add_run(line if line else ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(8.5)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(6)


def inject_code(doc, anchor_id):
    for title, code in CODE_INJECTIONS.get(anchor_id, []):
        add_code_block(doc, title, code)


def embed_image(doc, filename):
    path = IMG_MAP.get(filename)
    if path and path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(6)
        try:
            para.add_run().add_picture(str(path), width=Cm(14))
        except Exception as e:
            print(f"  [warn] could not embed {filename}: {e}")
    else:
        print(f"  [warn] image not found: {filename}")


INLINE_PAT = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')

def parse_inline(para, text):
    for part in INLINE_PAT.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); r.bold = True; r.font.name = 'Times New Roman'
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1]); r.italic = True; r.font.name = 'Times New Roman'
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1]); r.font.name = 'Courier New'; r.font.size = Pt(10)
        else:
            r = para.add_run(part); r.font.name = 'Times New Roman'


def add_normal(doc, text, justify=True):
    if not text.strip():
        return
    # Strip markdown link syntax [text](#anchor) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parse_inline(p, text)


def add_heading(doc, text, level):
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    h.paragraph_format.space_after  = Pt(4)


def process_table(doc, rows_lines):
    rows = []
    for line in rows_lines:
        if re.match(r'\|[-:\s|]+\|', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            txt = row[ci] if ci < len(row) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            parse_inline(para, txt)
            for run in para.runs:
                run.font.size = Pt(9); run.font.name = 'Times New Roman'
                if ri == 0: run.bold = True
    doc.add_paragraph()


def detect_imgs(line):
    return [f for f in IMG_MAP if f in line]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build():
    print(f"Reading {DRAFT}")
    raw = DRAFT.read_text(encoding='utf-8')

    # Apply text fixes
    for old, new in TEXT_FIXES:
        raw = raw.replace(old, new)

    # Append new supplementary references (if not already added)
    if '[S5]' not in raw:
        # Find the end of [S4] entry and append after
        raw = raw.rstrip() + '\n\n' + EXTRA_REFS.strip() + '\n'

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    lines = raw.split('\n')
    n     = len(lines)
    i     = 0

    while i < n:
        line = lines[i]

        # --- Anchor tag ---
        m = re.match(r'<a\s+id=["\']([^"\']+)["\']', line)
        if m:
            aid = m.group(1)
            prev = CODE_TRIGGER.get(aid)
            if prev:
                inject_code(doc, prev)
            i += 1
            continue

        # --- Separator ---
        if line.strip() == '---':
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            i += 1; continue

        # --- Skip other HTML tags (but process embed directives first) ---
        stripped = line.strip()
        m_embed = re.match(r'<!--\s*embed:(.+?)\s*-->', stripped)
        if m_embed:
            embed_image(doc, m_embed.group(1).strip())
            i += 1; continue
        if stripped.startswith('<') and stripped.endswith('>') and '<a ' not in stripped:
            i += 1; continue

        # --- Headings ---
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 3); i += 1; continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 2); i += 1; continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 1); i += 1; continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line[2:].strip())
            r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
            i += 1; continue

        # --- Fenced code block ---
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i]); i += 1
            i += 1
            for cl in code_lines:
                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(0)
                cp.paragraph_format.left_indent  = Cm(0.3)
                shade_paragraph(cp)
                r = cp.add_run(cl if cl else ' ')
                r.font.name = 'Courier New'; r.font.size = Pt(8.5)
            doc.add_paragraph()
            continue

        # --- Table ---
        if line.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].startswith('|'):
                tbl_lines.append(lines[i]); i += 1
            process_table(doc, tbl_lines)
            continue

        # --- Bullet list ---
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            parse_inline(p, line[2:].strip())
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(11)
            i += 1; continue

        # --- Numbered list ---
        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            parse_inline(p, re.sub(r'^\d+\.\s', '', line).strip())
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(11)
            i += 1; continue

        # --- Empty line ---
        if not stripped:
            i += 1; continue

        # --- Normal paragraph (may contain image refs) ---
        imgs = detect_imgs(line)
        add_normal(doc, line)
        for fname in imgs:
            if 'fig11_app_classify' not in fname:
                embed_image(doc, fname)
        i += 1

    print(f"Saving {OUT}")
    doc.save(str(OUT))
    size = OUT.stat().st_size / (1024*1024)
    print(f"Done — {size:.1f} MB")
    doc2 = Document(str(OUT))
    print(f"Paragraphs: {len(doc2.paragraphs)}, Tables: {len(doc2.tables)}")


if __name__ == '__main__':
    build()
